import os
import io
import re
import json
import tempfile
import unicodedata
from typing import List, Dict, Tuple

import streamlit as st
from groq import Groq
from PyPDF2 import PdfReader
from docx import Document as DocxReader
from fpdf import FPDF
from docx import Document as DocxWriter

# LangChain + embeddings + vectorstore
from langchain.text_splitters import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS

# ---------------------- Utilities ----------------------
def sanitize_for_latin(text: str) -> str:
    """
    Normalize text and remove/replace characters that cannot be encoded in latin-1.
    Ensures FPDF (latin-1) won't fail when writing.
    """
    if not isinstance(text, str):
        text = str(text)
    # Normalize unicode
    text = unicodedata.normalize("NFKD", text)
    # Replace smart quotes and common unicode dashes with ASCII counterparts
    replacements = {
        "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "-", "\u2026": "...", "\u00a0": " "
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    # Remove any remaining non-latin-1 characters by transliteration fallback
    encoded = text.encode("latin-1", errors="replace")
    return encoded.decode("latin-1")

def chunk_documents_from_uploads(uploaded_files) -> Tuple[List[str], List[dict]]:
    """
    Extract text from uploaded PDFs/DOCs and return text chunks + metadata list.
    metadata entries have keys: {'source': filename, 'page': page_number}
    """
    texts = []
    metadatas = []
    for file in uploaded_files:
        name = file.name
        temp_path = os.path.join(tempfile.gettempdir(), name)
        with open(temp_path, "wb") as f:
            f.write(file.read())
        if name.lower().endswith(".pdf"):
            try:
                reader = PdfReader(temp_path)
                for i, page in enumerate(reader.pages, start=1):
                    page_text = page.extract_text() or ""
                    texts.append(page_text)
                    metadatas.append({"source": name, "page": i})
            except Exception:
                # fallback: read entire file as raw
                with open(temp_path, "rb") as f:
                    raw = f.read().decode("utf-8", errors="ignore")
                    texts.append(raw)
                    metadatas.append({"source": name, "page": 0})
        elif name.lower().endswith(".docx"):
            try:
                doc = DocxReader(temp_path)
                full = "\n".join([p.text for p in doc.paragraphs])
                texts.append(full)
                metadatas.append({"source": name, "page": 0})
            except Exception:
                with open(temp_path, "rb") as f:
                    raw = f.read().decode("utf-8", errors="ignore")
                    texts.append(raw)
                    metadatas.append({"source": name, "page": 0})
        else:
            # plain text fallback
            try:
                with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
                    raw = f.read()
                texts.append(raw)
                metadatas.append({"source": name, "page": 0})
            except Exception:
                continue
    # Chunk documents into smaller passages for RAG
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = []
    chunk_metadatas = []
    for txt, meta in zip(texts, metadatas):
        pieces = splitter.split_text(txt or "")
        for idx, piece in enumerate(pieces):
            chunks.append(piece)
            # keep metadata useful: source + page + chunk_index
            chunk_metadatas.append({**meta, "chunk": idx})
    return chunks, chunk_metadatas

# ---------------------- RAG Index Builder ----------------------
def build_rag_index(chunks: List[str], metadatas: List[dict]) -> FAISS:
    """
    Build a FAISS vectorstore from list of text chunks using a HuggingFace embedding model.
    """
    embedder = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vect = FAISS.from_texts(texts=chunks, embedding=embedder, metadatas=metadatas)
    return vect

def retrieve_context(vectorstore: FAISS, by_pages: Tuple[int, int] = None, keywords: str = None, top_k: int = 6) -> str:
    """
    Retrieve relevant chunks by keywords or by page range. Returns concatenated context.
    """
    results = []
    if by_pages:
        start, end = by_pages
        # metadata search: linear filter across stored metadatas
        all_metas = vectorstore.index_to_docstore_id
        # FAISS vectorstore doesn't expose docstore easily; use similarity_search but filter by metadata
        candidates = vectorstore.similarity_search(keywords or "", k=top_k*5)
        for doc in candidates:
            meta = doc.metadata or {}
            page = meta.get("page", 0)
            if start <= page <= end:
                results.append(doc.page_content)
            if len(results) >= top_k:
                break
    elif keywords and keywords.strip():
        docs = vectorstore.similarity_search(keywords, k=top_k)
        results = [d.page_content for d in docs]
    else:
        # default: top k general chunks
        docs = vectorstore.similarity_search("", k=top_k)
        results = [d.page_content for d in docs]
    return "\n\n".join(results)

# ---------------------- Groq API Generator ----------------------
def generate_question_paper(content: str, specs: dict, with_answers: bool) -> str:
    """
    Calls Groq compound-mini model to generate a question paper.
    Returns model raw text or an error message string.
    """
    client = Groq(api_key=os.getenv("GROQ_API_KEY"))
    system_prompt = (
        "You are an expert academic exam paper generator. Produce a neatly formatted question paper "
        "with three sections: MCQs, Short Answer, Long Answer. Number questions sequentially per section. "
        "For MCQs, provide 4 options labeled (A)-(D). If with_answers is True include an 'Answers' section at the end with keys. "
        "If with_answers is False do NOT include any answers or solutions. Keep output plain text."
    )
    # Limit content length to avoid token explosion but include key retrieved context
    safe_content = content[:18000]  # truncated if huge
    user_prompt = (
        f"Study Material (context):\n{safe_content}\n\n"
        f"Requirements:\nTotal: {specs.get('total')}\nMCQ: {specs.get('mcq')}\n"
        f"Short: {specs.get('short')}\nLong: {specs.get('long')}\nInclude Answers: {with_answers}\n\n"
        "Produce the requested question paper. Ensure clarity and correct numbering. "
        "If material is insufficient, create sensible domain-appropriate questions based on the provided context."
    )
    try:
        resp = client.chat.completions.create(
            model="groq/compound-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
            max_tokens=4000,
        )
        return resp.choices[0].message.content
    except Exception as e:
        # Handle rate limit and return readable message
        err_str = str(e)
        if "429" in err_str or "rate" in err_str.lower():
            return "Error: Rate limit reached. Please try again later."
        return f"Error generating question paper: {err_str}"

# ---------------------- Save Utility ----------------------
def save_question_paper(text: str, filename: str, with_answers: bool) -> None:
    """
    Format text into clean layout and save as PDF and DOCX into outputs/ folder.
    If with_answers is False, strips sections named 'Answer' or 'Answers' and 'Explanation'.
    """
    os.makedirs("outputs", exist_ok=True)
    content = text or ""
    # Normalize spacing
    content = re.sub(r"\r\n", "\n", content)
    content = re.sub(r"\n\s+\n", "\n\n", content).strip()

    if not with_answers:
        # Try to remove 'Answers' sections heuristically
        content = re.split(r"(?i)\nanswers?\b", content)[0]

    # Save DOCX
    docx_path = os.path.join("outputs", f"{filename}.docx")
    doc = DocxWriter()
    doc.add_heading(filename.replace("_", " ").title(), level=1)
    for para in content.split("\n"):
        doc.add_paragraph(para)
    doc.save(docx_path)

    # Save PDF (sanitize for latin-1 to avoid encoding error)
    pdf_path = os.path.join("outputs", f"{filename}.pdf")
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVu", "", fname="", uni=True) if False else None  # placeholder for custom fonts
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=11)
    # sanitize line by line
    for line in content.split("\n"):
        safe_line = sanitize_for_latin(line)
        pdf.multi_cell(0, 7, safe_line)
    pdf.output(pdf_path)

# ---------------------- Streamlit App ----------------------
st.set_page_config(page_title="AI Question Paper Generator (RAG)", page_icon="📝")
st.title("📝 AI Question Paper Generator — RAG Edition")
st.caption("Upload notes, build a retrieval index, pick topics/pages, then generate Q-paper (Groq: compound-mini)")

# Upload
uploaded = st.file_uploader("Upload PDFs or DOCX textbooks/notes (multiple allowed)", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded:
    if st.button("🔎 Build RAG Index from uploads"):
        with st.spinner("Indexing documents... this may take a moment"):
            chunks, metas = chunk_documents_from_uploads(uploaded)
            if not chunks:
                st.error("No text extracted from uploads.")
            else:
                # Build FAISS index (in-memory)
                try:
                    vectorstore = build_rag_index(chunks, metas)
                    # Save vectorstore to session_state for reuse
                    st.session_state["vectorstore"] = vectorstore
                    st.session_state["chunks"] = chunks
                    st.session_state["metas"] = metas
                    st.success("✅ Index built. You can now retrieve by keywords or page ranges.")
                except Exception as e:
                    st.error(f"Index build error: {e}")

# Retrieval controls
st.subheader("Select context for question generation")
colA, colB = st.columns(2)
with colA:
    keywords = st.text_input("Search by topic keywords (e.g., 'photosynthesis chapter 3')", value="")
with colB:
    page_range = st.text_input("Or specify page range (start-end). Leave blank to ignore.", value="")

top_k = st.slider("Context size (number of chunks to retrieve)", 1, 12, 6)

# Question specs
st.subheader("Question Paper Specs")
total_q = st.number_input("Total Questions", min_value=1, value=20)
mcq_q = st.number_input("MCQs", min_value=0, value=10)
short_q = st.number_input("Short Answer", min_value=0, value=5)
long_q = st.number_input("Long Answer", min_value=0, value=2)
with_answers = st.checkbox("Include Answers (toggle)", value=True)
filename = st.text_input("Output filename (without extension)", value="question_paper")

if st.button("🚀 Generate Question Paper (RAG → Groq)"):
    if "vectorstore" not in st.session_state:
        st.warning("Please upload documents and build the RAG index first.")
    else:
        vectorstore: FAISS = st.session_state["vectorstore"]
        # Determine retrieval method
        pages_tuple = None
        if page_range and re.match(r"^\d+\s*-\s*\d+$", page_range.strip()):
            s, e = [int(x) for x in page_range.split("-")]
            pages_tuple = (min(s, e), max(s, e))
        with st.spinner("Retrieving context and generating question paper..."):
            try:
                context = retrieve_context(vectorstore, by_pages=pages_tuple, keywords=keywords, top_k=top_k)
                if not context.strip():
                    st.warning("No relevant context found — continuing with full uploaded text fallback.")
                    # fallback: concat all chunks
                    context = "\n\n".join(st.session_state.get("chunks", [])[:10])
                specs = {"total": total_q, "mcq": mcq_q, "short": short_q, "long": long_q}
                result = generate_question_paper(context, specs, with_answers)
                if result.startswith("Error"):
                    st.error(result)
                else:
                    save_question_paper(result, filename, with_answers)
                    st.success("✅ Generated and saved outputs in outputs/ directory")
                    # provide downloads
                    with open(f"outputs/{filename}.pdf", "rb") as fpdf_file:
                        st.download_button("📥 Download PDF", fpdf_file, file_name=f"{filename}.pdf")
                    with open(f"outputs/{filename}.docx", "rb") as fdocx:
                        st.download_button("📥 Download DOCX", fdocx, file_name=f"{filename}.docx")
            except Exception as e:
                st.error(f"❌ Error during generation: {e}")

st.info("Workflow: 1) Upload documents → 2) Build index → 3) Optionally search by keywords or page range → 4) Generate question paper.")
