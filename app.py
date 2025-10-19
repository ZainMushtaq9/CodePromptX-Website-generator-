# app.py
import os
import io
import re
import tempfile
import hashlib
import unicodedata
from typing import List, Tuple, Dict

import streamlit as st
from groq import Groq
from PyPDF2 import PdfReader
from docx import Document as DocxReader
from docx import Document as DocxWriter
from fpdf import FPDF
from PIL import Image
import fitz  # PyMuPDF

# LangChain & Embeddings (free, cloud-safe)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma

# Try Tesseract; if not available, fall back to Surya OCR (pure-Python)
OCR_ENGINE = None
try:
    import pytesseract

    # confirm tesseract binary exists
    try:
        _ = pytesseract.get_tesseract_version()
        OCR_ENGINE = "tesseract"
    except Exception:
        OCR_ENGINE = None
except Exception:
    OCR_ENGINE = None

if OCR_ENGINE is None:
    try:
        from surya_ocr import SuryaOCR

        surya = SuryaOCR()
        OCR_ENGINE = "surya"
    except Exception:
        OCR_ENGINE = None

# Groq client reads key from environment
GROQ_API_KEY = os.getenv("GROQ_API_KEY", None)
client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

st.set_page_config(page_title="AI Question Paper Generator (OCR+RAG)", page_icon="🧠", layout="wide")
st.title("🧠 AI Question Paper Generator — OCR + RAG")
st.caption("Supports text & scanned PDFs. Uses Groq (compound-mini), HuggingFace embeddings, Chroma. Caching enabled.")

# ------------------ Utilities ------------------
def file_hash(bytes_data: bytes) -> str:
    return hashlib.sha256(bytes_data).hexdigest()


def sanitize_for_latin(text: str) -> str:
    if not isinstance(text, str):
        text = str(text)
    text = unicodedata.normalize("NFKD", text)
    replacements = {
        "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "-", "\u2026": "...", "\u00a0": " "
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode("latin-1", errors="replace").decode("latin-1")


# ------------------ Extraction ------------------
@st.cache_data(show_spinner=False)
def extract_text_from_pdf_bytes(file_bytes: bytes) -> Tuple[List[str], List[Dict]]:
    """
    Extract text pages from a PDF file.
    Returns (list_of_page_texts, list_of_page_meta).
    Uses text extraction first; if page empty uses OCR engine (tesseract or surya).
    """
    pages_texts: List[str] = []
    metas: List[Dict] = []
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        path = tmp.name

    # Try text-based extraction per page
    try:
        reader = PdfReader(path)
        total_pages = len(reader.pages)
    except Exception:
        # If PyPDF2 fails, try opening with fitz for page count
        try:
            doc = fitz.open(path)
            total_pages = doc.page_count
            doc.close()
        except Exception:
            total_pages = 0

    # iterate pages and extract
    for p in range(total_pages):
        page_text = ""
        # Try PyPDF2 page extract
        try:
            reader = PdfReader(path)
            raw = reader.pages[p].extract_text() or ""
            if raw and len(raw.strip()) > 20:
                page_text = raw.strip()
        except Exception:
            page_text = ""

        # If page_text is empty, fallback to OCR using available engine
        if not page_text or len(page_text.strip()) < 20:
            ocr_text = ""
            try:
                doc = fitz.open(path)
                page = doc.load_page(p)
                pix = page.get_pixmap(dpi=200)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                if OCR_ENGINE == "tesseract":
                    ocr_text = pytesseract.image_to_string(img)
                elif OCR_ENGINE == "surya":
                    ocr_text = surya.extract_text(img)
                else:
                    ocr_text = ""
                doc.close()
            except Exception:
                ocr_text = ""

            if ocr_text and len(ocr_text.strip()) >= 20:
                page_text = ocr_text.strip()

        if page_text and len(page_text.strip()) >= 20:
            pages_texts.append(page_text)
            metas.append({"page": p + 1})
    # clean up tmp file
    try:
        os.remove(path)
    except Exception:
        pass
    return pages_texts, metas


@st.cache_data(show_spinner=False)
def extract_text_from_docx_bytes(file_bytes: bytes) -> Tuple[List[str], List[Dict]]:
    texts = []
    metas = []
    try:
        doc = DocxReader(io.BytesIO(file_bytes))
        full = "\n".join([p.text for p in doc.paragraphs if p.text and p.text.strip()])
        if full and len(full.strip()) > 20:
            texts.append(full.strip())
            metas.append({"page": 0})
    except Exception:
        pass
    return texts, metas


@st.cache_data(show_spinner=False)
def extract_text_from_upload(uploaded) -> Tuple[List[str], List[Dict]]:
    """
    Detect file type and extract text pages. Returns (chunks_per_page, metas).
    Caches by file content under Streamlit caching.
    """
    name = uploaded.name.lower()
    raw = uploaded.read()
    if name.endswith(".pdf"):
        pages, metas = extract_text_from_pdf_bytes(raw)
    elif name.endswith(".docx"):
        pages, metas = extract_text_from_docx_bytes(raw)
    else:
        pages, metas = [], []
    return pages, metas


# ------------------ Chunking & RAG ------------------
def chunk_texts(pages: List[str], metas: List[Dict]) -> Tuple[List[str], List[Dict]]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    all_chunks = []
    all_metas = []
    for page_text, meta in zip(pages, metas):
        pieces = splitter.split_text(page_text)
        for idx, piece in enumerate(pieces):
            if piece and len(piece.strip()) > 30:
                all_chunks.append(piece.strip())
                meta_copy = dict(meta)
                meta_copy["chunk"] = idx
                all_metas.append(meta_copy)
    return all_chunks, all_metas


@st.cache_data(show_spinner=False)
def build_chroma_index(chunks: List[str], metadatas: List[Dict]):
    if not chunks:
        raise ValueError("No text to index.")
    embedder = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vectordb = Chroma.from_texts(chunks, embedding=embedder, metadatas=metadatas)
    return vectordb


# ------------------ Groq Generator ------------------
def generate_question_paper(content: str, specs: dict, with_answers: bool) -> str:
    if client is None:
        return "Error: GROQ_API_KEY not configured in environment."
    system_prompt = (
        "You are an expert exam paper generator. Create a well-structured question paper based on provided material. "
        "Sections: MCQs (4 options A-D), Short Answer, Long Answer. Number each section clearly. "
        "If include_answers is True, add an 'Answers' section at the end with keys only."
    )
    safe_content = content[:16000]  # limit to keep tokens reasonable
    user_prompt = (
        f"Study material excerpt:\n{safe_content}\n\n"
        f"Specs: total={specs.get('total')}, mcq={specs.get('mcq')}, short={specs.get('short')}, long={specs.get('long')}\n"
        f"Include Answers: {with_answers}\n\n"
        "Produce the question paper in plain text."
    )
    try:
        resp = client.chat.completions.create(
            model="groq/compound-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
            max_tokens=4000,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"Error from Groq API: {e}"


# ------------------ Save Outputs ------------------
def save_outputs(text: str, filename: str, with_answers: bool) -> Tuple[str, str]:
    os.makedirs("outputs", exist_ok=True)
    content = text or ""
    if not with_answers:
        # heuristic: remove Answers section if present
        content = re.split(r"(?i)\nanswers?\b", content)[0]

    docx_path = os.path.join("outputs", f"{filename}.docx")
    pdf_path = os.path.join("outputs", f"{filename}.pdf")

    # DOCX
    doc = DocxWriter()
    doc.add_heading(filename.replace("_", " ").title(), level=1)
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(docx_path)

    # PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=11)
    for line in content.split("\n"):
        pdf.multi_cell(0, 7, sanitize_for_latin(line))
    pdf.output(pdf_path)

    return docx_path, pdf_path


# ------------------ Streamlit UI ------------------
st.sidebar.header("Settings")
st.sidebar.markdown("OCR Engine: **Tesseract** if available else **Surya OCR** (pure-Python).")
st.sidebar.markdown(f"Detected OCR engine: **{OCR_ENGINE or 'none (OCR disabled)'}**")

uploaded_files = st.file_uploader("Upload PDF/DOCX notes (multiple)", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    total_pages_extracted = 0
    all_chunks = []
    all_metas = []
    skipped_files = []
    for f in uploaded_files:
        st.info(f"Processing: {f.name}")
        pages, metas = extract_text_from_upload(f)
        if not pages:
            st.warning(f"No readable text extracted from {f.name}.")
            skipped_files.append(f.name)
            continue
        chunks, chunk_metas = chunk_texts(pages, metas)
        all_chunks.extend(chunks)
        all_metas.extend(chunk_metas)
        total_pages_extracted += len(pages)

    if not all_chunks:
        st.error("❌ No valid text extracted from uploads. Try a clearer scan or a text-based PDF.")
    else:
        st.success(f"✅ Extracted {total_pages_extracted} pages → {len(all_chunks)} chunks.")
        build_index = st.button("Build Knowledge Index (Chroma)")
        if build_index:
            with st.spinner("Building vector index..."):
                try:
                    vectordb = build_chroma_index(all_chunks, all_metas)
                    st.session_state["vectordb"] = vectordb
                    st.success("✅ Vector index built and cached in session.")
                except Exception as e:
                    st.error(f"Index build failed: {e}")

# Question paper UI
st.header("Generate Question Paper")
col1, col2 = st.columns(2)
with col1:
    total = st.number_input("Total questions", min_value=1, value=20)
    mcq = st.number_input("MCQs", min_value=0, value=10)
    short = st.number_input("Short answer questions", min_value=0, value=5)
    long = st.number_input("Long answer questions", min_value=0, value=2)
with col2:
    include_answers = st.checkbox("Include answers (add answer key)", value=True)
    filename = st.text_input("Output filename (no extension)", value="question_paper")
    topic_keywords = st.text_input("Optional topic/keyword filter (e.g., 'photosynthesis')", value="")

if st.button("Generate Paper"):
    if "vectordb" not in st.session_state:
        st.error("Please upload documents, extract text, and build the knowledge index first.")
    else:
        vectordb = st.session_state["vectordb"]
        with st.spinner("Retrieving context from index..."):
            try:
                if topic_keywords and topic_keywords.strip():
                    docs = vectordb.similarity_search(topic_keywords, k=8)
                else:
                    docs = vectordb.similarity_search("", k=8)
                context = "\n\n".join([d.page_content for d in docs])
            except Exception as e:
                st.error(f"Retrieval error: {e}")
                context = ""

        if not context.strip():
            st.error("No relevant context retrieved. Try broader keywords or rebuild the index.")
        else:
            specs = {"total": total, "mcq": mcq, "short": short, "long": long}
            with st.spinner("Generating question paper (Groq)..."):
                result = generate_question_paper(context, specs, include_answers)
            if result.lower().startswith("error") or result.startswith("Error"):
                st.error(result)
            else:
                docx_path, pdf_path = save_outputs(result, filename, include_answers)
                st.success("✅ Paper generated.")
                with open(pdf_path, "rb") as fp:
                    st.download_button("Download PDF", fp, file_name=f"{filename}.pdf")
                with open(docx_path, "rb") as fd:
                    st.download_button("Download DOCX", fd, file_name=f"{filename}.docx")

st.markdown("---")
st.markdown("Notes: If running on Streamlit Cloud, Tesseract binary may not be available — the app will try to use Surya OCR (pure-Python) if installed. For best OCR results run locally with Tesseract installed.")
